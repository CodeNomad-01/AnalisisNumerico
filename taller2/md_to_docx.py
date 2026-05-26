"""Convierte SUSTENTACION.md a SUSTENTACION.docx preservando el formato."""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

BASE = Path(__file__).parent
SRC = BASE / "SUSTENTACION.md"
DST = BASE / "SUSTENTACION.docx"


INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*)"        # **bold**
    r"|(`[^`]+`)"             # `code`
    r"|(\*[^*]+\*)"           # *italic*
)


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, fill_hex):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)


def set_paragraph_border(paragraph, side, color_hex="888888", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color_hex)
    p_bdr.append(border)


def add_runs(paragraph, text):
    """Tokeniza 'text' con negrita / inline-code / itálica y añade runs."""
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])
        bold_match, code_match, italic_match = match.groups()
        if bold_match:
            run = paragraph.add_run(bold_match[2:-2])
            run.bold = True
        elif code_match:
            run = paragraph.add_run(code_match[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif italic_match:
            run = paragraph.add_run(italic_match[1:-1])
            run.italic = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def is_table_row(line):
    return line.startswith("|") and line.endswith("|")


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and is_table_row(lines[i].strip()):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    # quitar fila separadora (|---|---|)
    cleaned = [r for r in rows if not all(re.fullmatch(r":?-+:?", c or "-") for c in r)]
    return cleaned, i


def convert(md_text, doc):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # bloques de código ```
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            set_paragraph_shading(p, "F4F4F4")
            set_paragraph_border(p, "left", "BBBBBB", "12")
            run = p.add_run("\n".join(buf))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        # bloques de cita >  (narración de exposición)
        if stripped.startswith(">"):
            quote_buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                content = lines[i].strip()[1:].lstrip()
                quote_buf.append(content)
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_shading(p, "FBF7E6")
            set_paragraph_border(p, "left", "C9A227", "18")
            for idx, qline in enumerate(quote_buf):
                if idx > 0:
                    p.add_run("\n")
                add_runs(p, qline)
            for run in p.runs:
                run.italic = True
            continue

        # encabezados
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        # separador horizontal
        if stripped == "---":
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # tabla
        if is_table_row(stripped):
            rows, new_i = parse_table(lines, i)
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(ncols):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        text = row[c_idx] if c_idx < len(row) else ""
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs(p, text)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
                            set_cell_shading(cell, "DCE6F1")
            i = new_i
            continue

        # listas con guion
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        # listas numeradas
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        # línea en blanco
        if stripped == "":
            i += 1
            continue

        # texto en itálica suelto del cierre del documento
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1


def main():
    md = SRC.read_text(encoding="utf-8")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    convert(md, doc)
    doc.save(DST)
    print(f"OK -> {DST}")


if __name__ == "__main__":
    main()
